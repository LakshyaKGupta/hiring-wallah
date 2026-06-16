import io
import fitz  # PyMuPDF
import pdfplumber
import logging
from docx import Document

logger = logging.getLogger("hiring_wallah.resume_parser")

def parse_resume(file_bytes: bytes, filename: str = "") -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else "pdf"
    if suffix == "txt":
        return parse_txt(file_bytes)
    if suffix == "docx":
        return parse_docx(file_bytes)
    return parse_pdf(file_bytes)


def parse_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return ""


def parse_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def parse_pdf(pdf_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes.
    Tries pdfplumber first, falls back to PyMuPDF (fitz) if pdfplumber fails or returns insufficient text.
    """
    text = ""
    
    # 1. Try pdfplumber
    try:
        logger.info("Attempting resume extraction with pdfplumber...")
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            extracted_pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(page_text)
            text = "\n".join(extracted_pages).strip()
            
            if len(text) > 100:
                logger.info(f"Successfully extracted {len(text)} characters using pdfplumber.")
                return text
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}. Falling back to PyMuPDF.")
        
    # 2. Fallback to PyMuPDF (fitz)
    try:
        logger.info("Attempting resume extraction with PyMuPDF...")
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        extracted_pages = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                extracted_pages.append(page_text)
        text = "\n".join(extracted_pages).strip()
        logger.info(f"Extracted {len(text)} characters using PyMuPDF fallback.")
        return text
    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {e}")
        return ""
